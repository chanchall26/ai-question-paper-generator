import os
import json
import re
import uuid
import io
from datetime import datetime
from flask import Blueprint, request, jsonify, render_template, current_app, send_file
from sqlalchemy.sql.expression import func
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, Table
from reportlab.lib.styles import getSampleStyleSheet
import google.generativeai as genai
from .models import Question, Paper, PaperQuestion
from . import db

main = Blueprint("main", __name__)

@main.route("/")
def index():
    return render_template("index.html")

@main.route("/api/generate", methods=["POST"])
def generate_paper():
    data = request.get_json()
    current_app.logger.info(f"Incoming /api/generate payload: {data}")

    subject = data.get("subject")
    class_ = data.get("class")
    board = data.get("schoolBoard")
    school = data.get("schoolName")
    qdist = data.get("questionDistribution", {})
    ddist = data.get("difficultyDistribution", {})
    exam_name = data.get("examName")
    questions = []

    try:
        prompt = f"""
        You are an experienced {board} school teacher. 
        Create a question paper for Class {class_}, Subject: {subject}.
        School: {school}

        Question type distribution: {qdist}
        Difficulty distribution: {ddist}

        Output only valid JSON array. Each item must have:
        - type (MCQ, Fill in the Blanks, Short Answer, Long Answer, Case Study, etc.)
        - question (text of the question)
        - options (for MCQ only: provide exactly 4 options as a JSON list)
        - marks (marks per question)
        - difficulty (Easy, Medium, Hard)
        - answer (correct answer or solution for the question; if MCQ, give the correct option letter like "A" or "B")
        """

        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(prompt)
        raw_text = response.text.strip()

        try:
            questions = json.loads(raw_text)
        except json.JSONDecodeError:
            json_match = re.search(r"```json\s*(.*?)```", raw_text, re.DOTALL)
            if json_match:
                raw_text = json_match.group(1)
                questions = json.loads(raw_text)
            else:
                raise ValueError("AI returned invalid JSON.")

        for q in questions:
            q["question_text"] = q.get("question")
            q["question_type"] = q.get("type")
            new_q = Question(
                school_name=school,
                board=board,
                class_=class_,
                subject=subject,
                question_type=q.get("type"),
                difficulty=q.get("difficulty"),
                question_text=q.get("question"),
                marks=q.get("marks"),
                answer=q.get("answer", "Not provided"),
                source="AI"
            )
            db.session.add(new_q)
            db.session.flush()
            q['id'] = new_q.id
        db.session.commit()

    except Exception as e:
        current_app.logger.error(f"AI generation failed: {e}. Falling back to DB.")
        db.session.rollback()
        all_questions = []
        for qtype, info in qdist.items():
            # info is now a dict like {'count': 3, 'marks': 1}
            num_to_fetch = int(info['count'])
            marks = int(info['marks'])
            if num_to_fetch > 0:
                db_questions = (
                    Question.query
                    .filter_by(subject=subject, class_=class_, question_type=qtype, marks=marks)
                    .order_by(func.rand())
                    .limit(num_to_fetch)
                    .all()
                )
                all_questions.extend([q.as_dict() for q in db_questions])
        questions = all_questions

    paper_id = str(uuid.uuid4())[:8]
    pdf_filename = f"paper_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    word_filename = f"{paper_id}.docx"
    answer_key_filename = f"answer_key_{paper_id}.txt"

    paper_entry = Paper(
        paper_id=paper_id,
        exam_name=exam_name,
        school_name=school,
        board=board,
        class_=class_,
        subject=subject,
        total_questions=len(questions),
        total_marks=sum(int(q.get("marks", 0)) for q in questions),
        pdf_path=f"/static/papers/{pdf_filename}",
        word_path=f"/static/papers/{word_filename}",
        answer_key_path=f"/static/papers/{answer_key_filename}"
    )
    db.session.add(paper_entry)
    db.session.commit()

    for q in questions:
        pq = PaperQuestion(
            paper_id=paper_id,
            question_id=q['id'],
            question_text=q.get("question_text"),
            type=q.get("question_type"),
            difficulty=q.get("difficulty"),
            marks=q.get("marks"),
            options=q.get("options", []),
            answer=q.get("answer", "Not provided")
        )
        db.session.add(pq)
    db.session.commit()

    papers_dir = os.path.join(current_app.root_path, "static", "papers")
    os.makedirs(papers_dir, exist_ok=True)
    json_path = os.path.join(papers_dir, f"{paper_id}.json")

    summary = {
        "total_questions": len(questions),
        "total_marks": sum(int(q.get("marks", 0)) for q in questions)
    }

    # ---- SORT QUESTIONS BY TYPE BEFORE RETURNING ----
    type_order = [
        "MCQ", "Multiple Choice",
        "Fill in the Blanks", "Fill",
        "Short Answer", "Short",
        "Long Answer", "Long",
        "Matching", "Match", "Match the Following",
        "Case Study", "Case"
    ]
    def get_type_order(q):
        t = (q.get("type") or q.get("question_type") or "").strip()
        try:
            return type_order.index(t)
        except ValueError:
            return len(type_order)
    questions_sorted = sorted(questions, key=get_type_order)
    # --------------------------------------------------

    with open(json_path, "w", encoding="utf-8") as f:
        json.dump({
            "paper_id": paper_id,
            "examName": exam_name,
            "schoolName": school,
            "schoolBoard": board,
            "class": class_,
            "subject": subject,
            "questions": questions_sorted,
            "summary": summary
        }, f, indent=2, ensure_ascii=False)

    # Generate PDF
    pdf_path = os.path.join(papers_dir, pdf_filename)
    c = canvas.Canvas(pdf_path, pagesize=A4)
    width, height = A4

    c.setFont("Helvetica-Bold", 16)
    c.drawCentredString(width/2, height - 50, f"{school}")
    c.setFont("Helvetica", 12)

    if exam_name:
        c.drawCentredString(width/2, height - 70, exam_name)
    else:
        c.drawCentredString(width/2, height - 70, f"{board} Board Examination")

    c.drawCentredString(width/2, height - 90, f"Class {class_} - {subject}")
    c.drawRightString(width - 50, height - 110, f"Date: {datetime.now().strftime('%d-%m-%Y')}")
    c.line(50, height - 120, width - 50, height - 120)

    # Group questions by type for sections
    sections = {
        "Multiple Choice": [],
        "Fill in the Blanks": [],
        "Short Answer": [],
        "Long Answer": [],
        "Matching": [],
        "Case Study": []
    }
    for q in questions_sorted:
        q_type = q.get("question_type", "").strip()
        if q_type in ["MCQ", "Multiple Choice"]:
            q_type = "Multiple Choice"
        elif q_type in ["Fill in the Blanks", "Fill"]:
            q_type = "Fill in the Blanks"
        elif q_type in ["Short", "Short Answer"]:
            q_type = "Short Answer"
        elif q_type in ["Long", "Long Answer"]:
            q_type = "Long Answer"
        elif q_type in ["Matching", "Match", "Match the Following"]:
            q_type = "Matching"
        elif q_type in ["Case Study", "Case"]:
            q_type = "Case Study"
        if q_type in sections:
            sections[q_type].append(q)

    section_titles = {
        "Multiple Choice": "Section A - Multiple Choice Questions",
        "Fill in the Blanks": "Section B - Fill in the Blanks",
        "Short Answer": "Section C - Short Answer Questions",
        "Long Answer": "Section D - Long Answer Questions",
        "Matching": "Section E - Matching Questions",
        "Case Study": "Section F - Case Study"
    }

    y = height - 150
    qnum = 1
    styles = getSampleStyleSheet()
    styleN = styles["Normal"]

    for sec, qlist in sections.items():
        if not qlist:
            continue
        c.setFont("Helvetica-Bold", 13)
        c.drawString(50, y, section_titles[sec])
        y -= 25
        for q in qlist:
            text = f"Q{qnum}. {q.get('question_text')}   ({q.get('marks', 0)} marks)"
            p = Paragraph(text, styleN)
            w, h = p.wrap(width - 100, y)
            if y - h < 100:
                c.showPage()
                y = height - 50
            p.drawOn(c, 70, y - h)
            y -= (h + 15)
            if sec == "Multiple Choice":
                options = q.get("options", [])
                for idx, opt in enumerate(options, start=1):
                    opt_text = f"   ({chr(64+idx)}) {opt}"
                    p_opt = Paragraph(opt_text, styleN)
                    w, h = p_opt.wrap(width - 120, y)
                    if y - h < 100:
                        c.showPage()
                        y = height - 50
                    p_opt.drawOn(c, 90, y - h)
                    y -= (h + 10)
            qnum += 1
        y -= 15
    c.save()

    return jsonify({
        "questions": [{
            "id": q.get("id"),
            "question_text": q.get("question_text"),
            "marks": q.get("marks"),
            "difficulty": q.get("difficulty"),
            "type": q.get("type") or q.get("question_type"),
            "source": q.get("source", "Database")
        } for q in questions_sorted],
        "summary": summary,
        "pdf_url": f"/static/papers/{pdf_filename}",
        "word_url": f"/api/download/word/{paper_id}",
        "answer_key_url": f"/api/download/answer_key/{paper_id}"
    })

@main.route("/api/download/word/<paper_id>", methods=["GET"])
def download_word(paper_id):
    json_path = os.path.join(current_app.root_path, "static", "papers", f"{paper_id}.json")
    if not os.path.exists(json_path):
        return jsonify({"error": "Paper not found"}), 404
    with open(json_path, "r", encoding="utf-8") as f:
        paper = json.load(f)
    doc = Document()
    doc.add_heading(paper.get("examName", "Question Paper"), 0)
    doc.add_paragraph(f"School: {paper.get('schoolName')}")
    doc.add_paragraph(f"Board: {paper.get('schoolBoard')}")
    doc.add_paragraph(f"Class: {paper.get('class')}  Subject: {paper.get('subject')}")
    doc.add_heading("Questions", level=1)
    for i, q in enumerate(paper.get("questions", []), 1):
        doc.add_paragraph(f"Q{i}. {q['question_text']} ({q['marks']} marks) [{q['difficulty']}]")
        if q.get("question_type") in ["MCQ", "Multiple Choice"]:
            options = q.get("options", [])
            for idx, opt in enumerate(options, start=1):
                doc.add_paragraph(f"   ({chr(64+idx)}) {opt}", style="List Bullet")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name=f"paper_{paper_id}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@main.route("/api/download/answer_key/<paper_id>", methods=["GET"])
def download_answer_key(paper_id):
    json_path = os.path.join(current_app.root_path, "static", "papers", f"{paper_id}.json")
    if not os.path.exists(json_path):
        return jsonify({"error": "Paper not found"}), 404
    with open(json_path, "r", encoding="utf-8") as f:
        paper = json.load(f)
    answer_text = f"Answer Key for {paper.get('subject')} - Class {paper.get('class')}\n\n"
    for i, q in enumerate(paper.get("questions", []), 1):
        answer_text += f"Answer {i}. {q.get('answer', 'Answer not available')}\n\n"
    buf = io.BytesIO(answer_text.encode("utf-8"))
    return send_file(
        buf,
        as_attachment=True,
        download_name=f"answer_key_{paper_id}.txt",
        mimetype="text/plain"
    )
